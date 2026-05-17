import os
import re
from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

BASE_DIR = r"c:\Users\dell\Desktop\科创中心\大创\ZaRan-main\ZaRan-main\ZaRan"
CLIENT_DIR = r"c:\Users\dell\Desktop\科创中心\大创\ZaRan-main\ZaRan-main\ZaRan\ClientApp"

BACKEND_FILES = [
    "Abstraction\\Models\\DbModels\\User.cs",
    "Abstraction\\Models\\DbModels\\Article.cs",
    "Abstraction\\Models\\DbModels\\Good.cs",
    "Abstraction\\Models\\DbModels\\Collection.cs",
    "Abstraction\\Models\\DbModels\\Collect.cs",
    "Abstraction\\Models\\DbModels\\Comment.cs",
    "Abstraction\\Models\\DbModels\\Message.cs",
    "Abstraction\\Models\\DbModels\\RelationShip.cs",
    "Abstraction\\Models\\DbModels\\Transaction.cs",
    "Abstraction\\Models\\Dtos\\ArrayResult.cs",
    "Abstraction\\Models\\Dtos\\PaginationRequest.cs",
    "Extensions\\AuthorizationExtension\\UserExtensions.cs",
    "Services\\AppDbContext.cs",
    "Endpoints\\UsersEndpoints\\UserRegisterEndpoint.cs",
    "Endpoints\\UsersEndpoints\\UserLoginEndpoint.cs",
    "Endpoints\\UsersEndpoints\\UserPatchEndpoint.cs",
    "Endpoints\\UsersEndpoints\\UserListEndpoint.cs",
    "Endpoints\\UsersEndpoints\\UserDetailEndpoint.cs",
    "Endpoints\\UsersEndpoints\\UserStatusEndpoint.cs",
    "Endpoints\\UsersEndpoints\\UserAvatarUploadEndpoint.cs",
    "Endpoints\\ArticlesEndpoints\\ArticlePostEndpoint.cs",
    "Endpoints\\ArticlesEndpoints\\ArticlePatchEndpoint.cs",
    "Endpoints\\ArticlesEndpoints\\ArticleDeleteEndpoint.cs",
    "Endpoints\\ArticlesEndpoints\\ArticleGetEndpoint.cs",
    "Endpoints\\ArticlesEndpoints\\ArticleListEndpoint.cs",
    "Endpoints\\ArticlesEndpoints\\ArticleRejectEndpoint.cs",
    "Endpoints\\GoodsEndpoints\\GoodPostEndpoint.cs",
    "Endpoints\\GoodsEndpoints\\GoodPatchEndpoint.cs",
    "Endpoints\\GoodsEndpoints\\GoodDeleteEndpoint.cs",
    "Endpoints\\GoodsEndpoints\\GoodGetEndpoint.cs",
    "Endpoints\\GoodsEndpoints\\GoodListEndpoint.cs",
    "Endpoints\\GoodsEndpoints\\GoodRejectEndpoint.cs",
    "Endpoints\\GoodsEndpoints\\GoodCollectsListEndpoint.cs",
    "Endpoints\\CollectionsEndpoints\\CollectionPostEndpoint.cs",
    "Endpoints\\CollectionsEndpoints\\CollectionPatchEndpoint.cs",
    "Endpoints\\CollectionsEndpoints\\CollectionDeleteEndpoint.cs",
    "Endpoints\\CollectionsEndpoints\\CollectionGetEndpoint.cs",
    "Endpoints\\CollectionsEndpoints\\CollectionListEndpoint.cs",
    "Endpoints\\CollectionsEndpoints\\CollectionItemsEndpoint.cs",
    "Endpoints\\CommentsEndpoints\\CommentPostEndpoint.cs",
    "Endpoints\\CommentsEndpoints\\CommentListEndpoint.cs",
    "Endpoints\\CommentsEndpoints\\CommentGetEndpoint.cs",
    "Endpoints\\MessagesEndpoints\\MessagePostEndpoint.cs",
    "Endpoints\\MessagesEndpoints\\MessageListEndpoint.cs",
    "Endpoints\\MessagesEndpoints\\MessageGetEndpoint.cs",
    "Endpoints\\TransactionsEndpoints\\TransactionPostEndpoint.cs",
    "Endpoints\\TransactionsEndpoints\\TransactionPatchEndpoint.cs",
    "Endpoints\\TransactionsEndpoints\\TransactionListEndpoint.cs",
    "Endpoints\\TransactionsEndpoints\\TransactionGetEndpoint.cs",
    "Endpoints\\RelationshipsEndpoints\\UserRelationshipAddEndpoint.cs",
    "Endpoints\\RelationshipsEndpoints\\UserRelationshipDeleteEndpoint.cs",
    "Endpoints\\RelationshipsEndpoints\\UserRelationshipGetEndpoint.cs",
    "Endpoints\\CommonEndpoints\\StaticUploadEndpoint.cs",
    "Endpoints\\CommonEndpoints\\StaticGetEndpoint.cs",
    "Endpoints\\CommonEndpoints\\SeedingEndpoint.cs",
    "Endpoints\\CommentsEndpoints\\CommentPatchEndpoint.cs",
    "Endpoints\\CommentsEndpoints\\CommentDeleteEndpoint.cs",
    "Endpoints\\MessagesEndpoints\\MessageDeleteEndpoint.cs",
    "Endpoints\\MessagesEndpoints\\MessageReadEndpoint.cs",
    "Endpoints\\MessagesEndpoints\\MessageReadAllEndpoint.cs",
    "Endpoints\\TransactionsEndpoints\\TransactionDeleteEndpoint.cs",
]

FRONTEND_FILES = [
    "composables\\useApi.ts",
    "composables\\useArticles.ts",
    "composables\\useGoods.ts",
    "composables\\useTransactions.ts",
    "composables\\useMessages.ts",
    "composables\\useComments.ts",
    "composables\\useRelationships.ts",
    "composables\\useCollections.ts",
    "composables\\useArticleManagement.ts",
    "composables\\useCommunity.ts",
    "composables\\useGeneLibrary.ts",
    "composables\\useNotifications.ts",
    "composables\\usePermissions.ts",
    "composables\\useLoading.ts",
]

def remove_comments_cs(lines):
    result = []
    in_multiline = False
    for line in lines:
        stripped = line.strip()
        if in_multiline:
            if '*/' in stripped:
                in_multiline = False
                after = stripped[stripped.index('*/') + 2:]
                if after.strip():
                    result.append(after if not line.startswith(' ') else line[line.index('*/')+2:])
            continue
        if stripped.startswith('/*'):
            if '*/' in stripped:
                after = stripped[stripped.index('*/') + 2:]
                if after.strip():
                    result.append(after if not line.startswith(' ') else line[line.index('*/')+2:])
                continue
            in_multiline = True
            continue
        if stripped.startswith('///'):
            continue
        comment_idx = None
        in_str = False
        in_char = False
        i = 0
        while i < len(stripped):
            c = stripped[i]
            if c == '"' and not in_char:
                in_str = not in_str
            elif c == "'" and not in_str:
                in_char = not in_char
            elif c == '/' and not in_str and not in_char:
                if i + 1 < len(stripped) and stripped[i+1] == '/':
                    comment_idx = i
                    break
            i += 1
        if comment_idx is not None:
            if comment_idx > 0:
                result.append(stripped[:comment_idx].rstrip())
        else:
            result.append(stripped)
    return result

def remove_comments_ts(lines):
    result = []
    in_multiline = False
    for line in lines:
        stripped = line.strip()
        if in_multiline:
            if '*/' in stripped:
                in_multiline = False
                after = stripped[stripped.index('*/') + 2:]
                if after.strip():
                    result.append(after)
            continue
        if stripped.startswith('/*') or stripped.startswith('/**'):
            if '*/' in stripped:
                after = stripped[stripped.index('*/') + 2:]
                if after.strip():
                    result.append(after)
                continue
            in_multiline = True
            continue
        if stripped.startswith('*') and not stripped.startswith('* '):
            continue
        if stripped.startswith('* '):
            s = stripped[1:].strip()
            if s:
                result.append(s)
            continue
        comment_idx = None
        in_str = False
        in_char = False
        in_template = False
        i = 0
        while i < len(stripped):
            c = stripped[i]
            if c == '`' and not in_str:
                in_template = not in_template
            elif c == '"' and not in_char and not in_template:
                in_str = not in_str
            elif c == "'" and not in_str and not in_template:
                in_char = not in_char
            elif c == '/' and not in_str and not in_char and not in_template:
                if i + 1 < len(stripped) and stripped[i+1] == '/':
                    comment_idx = i
                    break
            i += 1
        if comment_idx is not None:
            if comment_idx > 0:
                result.append(stripped[:comment_idx].rstrip())
        else:
            result.append(stripped)
    return result

def process_file(filepath, is_cs=True):
    if not os.path.exists(filepath):
        return []
    with open(filepath, 'r', encoding='utf-8') as f:
        lines = f.readlines()
    lines = [l.rstrip('\n\r') for l in lines]
    if is_cs:
        lines = remove_comments_cs(lines)
    else:
        lines = remove_comments_ts(lines)
    result = [l for l in lines if l.strip() != '']
    return result

def set_cell_font(cell, font_name='宋体', font_size=Pt(9)):
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.font.name = font_name
            run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
            run.font.size = font_size

def add_code_line(doc, text, line_num):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.0
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    num_run = p.add_run(f"{line_num:4d}  ")
    num_run.font.name = '宋体'
    num_run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    num_run.font.size = Pt(9)
    num_run.font.color.rgb = RGBColor(128, 128, 128)
    code_run = p.add_run(text)
    code_run.font.name = '宋体'
    code_run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    code_run.font.size = Pt(9)

def generate_docx():
    doc = Document()
    
    style = doc.styles['Normal']
    font = style.font
    font.name = '宋体'
    font.size = Pt(9)
    style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    
    for section in doc.sections:
        section.page_width = Cm(21.0)
        section.page_height = Cm(29.7)
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.0)
        section.right_margin = Cm(2.0)
    
    all_lines = []
    
    header = f"软件名称：檐下风铃-扎染文化数字平台 V1.0"
    all_lines.append(header)
    all_lines.append(f"开发者：郑骐嵘（个人）")
    all_lines.append(f"开发完成日期：2026年3月1日")
    all_lines.append(f"开发语言：C# (.NET 9.0) + TypeScript (Vue 3 / Nuxt 3)")
    all_lines.append(f"后端框架：FastEndpoints + Entity Framework Core")
    all_lines.append(f"前端框架：Nuxt.js 3 + Vue 3 Composition API")
    all_lines.append("")
    
    for rel_path in BACKEND_FILES:
        full_path = os.path.join(BASE_DIR, rel_path)
        if os.path.exists(full_path):
            lines = process_file(full_path, is_cs=True)
            if lines:
                all_lines.append(f"=== {rel_path} ===")
                all_lines.extend(lines)
    
    for rel_path in FRONTEND_FILES:
        full_path = os.path.join(CLIENT_DIR, rel_path)
        if os.path.exists(full_path):
            lines = process_file(full_path, is_cs=False)
            if lines:
                all_lines.append(f"=== {rel_path} ===")
                all_lines.extend(lines)

    program_path = os.path.join(BASE_DIR, "Program.cs")
    if os.path.exists(program_path):
        prog_lines = process_file(program_path, is_cs=True)
        if prog_lines:
            all_lines.append("=== Program.cs ===")
            all_lines.extend(prog_lines)
    
    lines_per_page = 50
    page_num = 1
    total_lines = len(all_lines)
    total_pages = (total_lines + lines_per_page - 1) // lines_per_page
    
    print(f"Total lines: {total_lines}")
    print(f"Total pages: {total_pages}")
    
    front_30_pages = min(30, total_pages)
    back_30_start = max(0, total_pages - 30)
    
    pages_to_include = list(range(front_30_pages)) + list(range(back_30_start, total_pages))
    pages_to_include = sorted(set(pages_to_include))
    
    for pg in pages_to_include:
        start_line = pg * lines_per_page
        end_line = min(start_line + lines_per_page, total_lines)
        
        for i, line in enumerate(all_lines[start_line:end_line]):
            global_line_num = start_line + i + 1
            add_code_line(doc, line, global_line_num)
        
        if pg < total_pages - 1 or pg == pages_to_include[-1]:
            doc.add_page_break()
    
    output_path = os.path.join(os.path.dirname(BASE_DIR), "檐下风铃-扎染文化数字平台_V1.0_源代码文档.docx")
    doc.save(output_path)
    print(f"Document saved to: {output_path}")
    print(f"Pages generated: {len(pages_to_include)}")
    return output_path, total_pages, total_lines

if __name__ == "__main__":
    generate_docx()