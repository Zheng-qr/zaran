from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.enum.section import WD_SECTION
from docx.shared import RGBColor

def create_heading(doc, text, level):
    p = doc.add_heading(text, level)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    return p

def add_paragraph(doc, text, font_size=12, bold=False, align=WD_ALIGN_PARAGRAPH.LEFT):
    p = doc.add_paragraph()
    p.alignment = align
    run = p.add_run(text)
    run.font.size = Pt(font_size)
    run.font.name = '宋体'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    run.bold = bold
    return p

def add_table(doc, data, headers):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'
    
    hdr_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        hdr_cells[i].text = header
    
    for row in data:
        row_cells = table.add_row().cells
        for i, cell in enumerate(row):
            row_cells[i].text = cell
    
    return table

def generate_manual():
    doc = Document()
    
    style = doc.styles['Normal']
    font = style.font
    font.name = '宋体'
    font.size = Pt(12)
    style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    
    for section in doc.sections:
        section.page_width = Cm(21.0)
        section.page_height = Cm(29.7)
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.0)
        section.right_margin = Cm(2.0)
    
    doc.add_heading('计算机软件著作权登记软件 V1.0 说明书', 0)
    add_paragraph(doc, '【软件全称】檐下风铃-扎染文化数字平台 V1.0 软件说明书', 12, bold=True)
    
    create_heading(doc, '第 1 章 引言', 1)
    
    create_heading(doc, '1.1 编写目的', 2)
    add_paragraph(doc, '本文档用于【檐下风铃-扎染文化数字平台】V1.0 的计算机软件著作权登记申请，详细描述该软件的功能特点、操作流程、运行环境及使用方法，供软件著作权登记机构审查使用。')
    
    create_heading(doc, '1.2 软件基本信息', 2)
    add_paragraph(doc, '软件全称：檐下风铃-扎染文化数字平台', 12)
    add_paragraph(doc, '软件简称：檐下风铃', 12)
    add_paragraph(doc, '版本号：V1.0', 12)
    add_paragraph(doc, '开发完成日期：2026年03月01日', 12)
    add_paragraph(doc, '首次发表日期：未发表', 12)
    add_paragraph(doc, '开发者：郑骐嵘（个人）', 12)
    add_paragraph(doc, '权利归属：郑骐嵘（个人）', 12)
    
    create_heading(doc, '1.3 术语定义', 2)
    add_paragraph(doc, '无专用术语需要定义。')
    
    create_heading(doc, '1.4 参考资料', 2)
    add_paragraph(doc, '【项目需求说明书 V1.0】', 12)
    add_paragraph(doc, '【系统设计说明书 V1.0】', 12)
    add_paragraph(doc, '【软件测试报告 V1.0】', 12)
    add_paragraph(doc, '（图片预留位置）', 12)
    
    create_heading(doc, '第 2 章 软件概述', 1)
    
    create_heading(doc, '2.1 软件功能概述', 2)
    add_paragraph(doc, '本软件是一款面向内容创作者和消费者的综合平台用户的内容创作与数字商品交易管理平台，主要解决自贡扎染及同类非遗项目面临的五大行业痛点：', 12)
    add_paragraph(doc, '1. 认知壁垒厚重：64.6%受访者未听说过扎染，大众普遍存在"云南=扎染"的刻板印象，自贡扎染独特的蜀地技艺与文化价值被严重低估，市场辨识度不足。', 12)
    add_paragraph(doc, '2. 传承断代危机：仅7位受认可的自贡扎染传承人且普遍老龄化，年轻群体因收入低、学习周期长参与意愿低，传统"口传心授"模式易导致技艺流失。', 12)
    add_paragraph(doc, '3. 商业成本高企：纯手工制作需经历数十道工序，单件耗时72小时以上，天然原料与人工成本导致产品定价偏高，传统品类单一缺乏高溢价变现渠道。', 12)
    add_paragraph(doc, '4. 版权保护缺失：传统扎染纹样易被抄袭盗用，缺乏数字化确权与溯源机制，非遗传承人与设计师的合法权益难以保障。', 12)
    add_paragraph(doc, '5. 传播形式滞后：固守传统表达形式，与Z世代审美脱节，难以吸引年轻群体关注与消费，非遗文化陷入"博物馆式保护"的困境。', 12)
    add_paragraph(doc, '平台打造8大核心功能模块，形成"内容-交易-共创-服务"的完整闭环：', 12)
    add_paragraph(doc, '【首页】提供IP故事、科普百科、基因库等板块的快速导航；设置横幅展览位，出售轮播广告权提升版权方曝光；长期保留自贡扎染IP专属展示位', 12)
    add_paragraph(doc, '【IP故事页】展示入驻IP的宣传视频、原创角色（OC）设定、背景世界观、人物关系图谱；以自贡扎染为例，详细呈现"盐夏""绞缬"等染灵系列的完整设定', 12)
    add_paragraph(doc, '【科普百科】以文章、短视频形式科普非遗文化，讲解扎染工艺流程、地域特色差异（如自贡热染与云南冷染的区别）；发布传承人访谈与技艺教学内容', 12)
    add_paragraph(doc, '【纹样基因库】存储经区块链确权的非遗核心纹样，分为免费开放与版权授权两类；支持按主题、风格、地域检索纹样；点击可跳转至版权所有者页面完成交易，支持第三方非遗基因库嫁接', 12)
    add_paragraph(doc, '【创意工坊】提供在线设计白板与基因库素材，支持用户组合创新设计；涉及版权素材需购买授权后方可导出成果；支持发布设计"召集令"，面向平台用户招标征集作品', 12)
    add_paragraph(doc, '【文创商城】分类展示版权作品与实体文创商品，支持按关键词、价格、品类搜索；提供智能合约自动交易与买卖双方直接沟通两种模式；基于用户行为数据进行个性化推荐', 12)
    add_paragraph(doc, '【共创社区】设置"我卖版权""我要版权""分享"三大发帖板块，平台根据用户需求精准推流；定期发布官方共创活动，出售活动置顶权提升第三方IP曝光', 12)
    add_paragraph(doc, '【个人中心】管理账户流水、交易记录、发布/已购/已售作品、收藏关注；支持购买作品额外流量、设计3D虚拟形象并装扮；提供账号设置、消息通知、非遗传承人资质查询等基础功能', 12)
    add_paragraph(doc, '平台核心价值包括：', 12)
    add_paragraph(doc, '1. 文化传承价值：数字化永久保存非遗技艺，年轻化破圈传播', 12)
    add_paragraph(doc, '2. 商业赋能价值：多元变现体系，柔性生产模式', 12)
    add_paragraph(doc, '3. 社会公益价值：赋能乡村振兴，增强文化认同', 12)
    add_paragraph(doc, '4. 行业示范价值：可复制非遗活化范式', 12)
    
    create_heading(doc, '2.2 软件运行环境', 2)
    
    create_heading(doc, '2.2.1 服务器端环境', 3)
    add_paragraph(doc, '操作系统：Windows Server 2022 / Ubuntu 22.04 LTS 64位', 12)
    add_paragraph(doc, 'CPU：Intel Xeon E5-2670 v3 及以上', 12)
    add_paragraph(doc, '内存：8GB 及以上', 12)
    add_paragraph(doc, '硬盘：100GB 及以上可用空间', 12)
    add_paragraph(doc, '数据库：PostgreSQL 15+ / Entity Framework Core (开发环境使用 InMemory)', 12)
    add_paragraph(doc, '开发语言：C# .NET 9.0', 12)
    add_paragraph(doc, '后端框架：FastEndpoints', 12)
    add_paragraph(doc, 'Web服务器：Kestrel', 12)
    
    create_heading(doc, '2.2.2 客户端环境', 3)
    add_paragraph(doc, '操作系统：Windows 10/11 64位、macOS 12及以上、Linux主流发行版', 12)
    add_paragraph(doc, '浏览器：Google Chrome 100及以上、Microsoft Edge 100及以上、Mozilla Firefox 100及以上', 12)
    add_paragraph(doc, '分辨率：1920×1080及以上', 12)
    
    create_heading(doc, '2.3 软件整体架构', 2)
    add_paragraph(doc, '本软件采用经典的三层B/S架构设计，分为表现层、业务逻辑层和数据访问层：', 12)
    add_paragraph(doc, '表现层：负责与用户交互，展示系统界面和数据，使用HTML、CSS、JavaScript（Vue 3 / Nuxt 3）开发', 12)
    add_paragraph(doc, '业务逻辑层：实现系统的核心业务逻辑，处理用户请求，使用C# .NET 9.0 + FastEndpoints框架开发', 12)
    add_paragraph(doc, '数据访问层：负责与数据库交互，实现数据的增删改查操作，使用Entity Framework Core + PostgreSQL', 12)
    add_paragraph(doc, '【此处插入软件架构图】', 12)
    
    create_heading(doc, '第 3 章 软件安装与部署', 1)
    
    create_heading(doc, '3.1 服务器端部署步骤', 2)
    add_paragraph(doc, '1. 安装操作系统（Ubuntu 22.04 LTS），配置网络环境', 12)
    add_paragraph(doc, '2. 安装.NET 9.0 SDK、PostgreSQL 15和Nginx', 12)
    add_paragraph(doc, '3. 克隆项目代码到服务器指定目录：git clone【项目地址】', 12)
    add_paragraph(doc, '4. 安装项目依赖：dotnet restore', 12)
    add_paragraph(doc, '5. 配置数据库连接信息，更新appsettings.json', 12)
    add_paragraph(doc, '6. 构建项目：dotnet build --configuration Release', 12)
    add_paragraph(doc, '7. 启动服务：dotnet run --configuration Release', 12)
    add_paragraph(doc, '8. 配置Nginx反向代理，设置域名解析', 12)
    add_paragraph(doc, '9. 访问系统地址，验证部署成功', 12)
    
    create_heading(doc, '3.2 客户端访问方式', 2)
    add_paragraph(doc, '1. 打开浏览器，输入系统访问地址：【例如：http://www.yanxiafengling.com】', 12)
    add_paragraph(doc, '2. 进入系统登录界面，输入账号密码即可使用', 12)
    
    create_heading(doc, '第 4 章 功能模块详细说明', 1)
    
    create_heading(doc, '4.1 用户登录模块', 2)
    create_heading(doc, '4.1.1 功能描述', 3)
    add_paragraph(doc, '实现用户身份验证，支持账号密码登录。登录成功后，系统根据用户角色自动跳转至对应的主界面；登录失败时，给出明确的错误提示。', 12)
    create_heading(doc, '4.1.2 操作步骤', 3)
    add_paragraph(doc, '1. 打开浏览器，输入系统访问地址，进入登录界面', 12)
    add_paragraph(doc, '2. 输入账号和密码', 12)
    add_paragraph(doc, '3. 点击"登录"按钮', 12)
    add_paragraph(doc, '4. 验证通过后，跳转至对应角色的主界面；验证失败，提示"账号或密码错误"', 12)
    create_heading(doc, '4.1.3 运行截图', 3)
    add_paragraph(doc, '【此处插入登录界面截图】', 12)
    
    create_heading(doc, '4.2 用户管理模块', 2)
    create_heading(doc, '4.2.1 功能描述', 3)
    add_paragraph(doc, '管理员可对系统用户进行管理，包括查看用户列表、查看用户详情、修改用户信息、重置用户密码、分配用户权限等操作。', 12)
    create_heading(doc, '4.2.2 操作步骤', 3)
    add_paragraph(doc, '1. 管理员登录系统，进入"用户管理"页面', 12)
    add_paragraph(doc, '2. 查看用户列表，可按关键词搜索用户', 12)
    add_paragraph(doc, '3. 点击用户列表中的"详情"按钮，查看用户详细信息', 12)
    add_paragraph(doc, '4. 点击用户列表中的"编辑"按钮，可修改用户信息', 12)
    add_paragraph(doc, '5. 点击用户列表中的"重置密码"按钮，可重置用户密码', 12)
    create_heading(doc, '4.2.3 运行截图', 3)
    add_paragraph(doc, '【此处插入用户管理界面截图】', 12)
    
    create_heading(doc, '4.3 数字版权交易业务', 2)
    create_heading(doc, '4.3.1 功能描述', 3)
    add_paragraph(doc, '提供非遗纹样版权的交易服务，支持版权作品的发布、购买、授权等操作。用户可上传原创设计作品，设置版权授权方式，其他用户可购买授权后使用。', 12)
    create_heading(doc, '4.3.2 操作步骤', 3)
    add_paragraph(doc, '1. 用户登录系统，进入"文创商城"或"共创社区"', 12)
    add_paragraph(doc, '2. 选择"我卖版权"板块，发布版权作品', 12)
    add_paragraph(doc, '3. 设置作品名称、描述、授权方式、价格等信息', 12)
    add_paragraph(doc, '4. 其他用户浏览版权作品，选择购买授权', 12)
    add_paragraph(doc, '5. 完成交易后，买家获得作品使用权', 12)
    create_heading(doc, '4.3.3 运行截图', 3)
    add_paragraph(doc, '【此处插入版权交易界面截图】', 12)
    
    create_heading(doc, '4.4 非遗基因库服务业务', 2)
    create_heading(doc, '4.4.1 功能描述', 3)
    add_paragraph(doc, '存储经区块链确权的非遗核心纹样，分为免费开放与版权授权两类。支持按主题、风格、地域检索纹样，点击可跳转至版权所有者页面完成交易，支持第三方非遗基因库嫁接。', 12)
    create_heading(doc, '4.4.2 操作步骤', 3)
    add_paragraph(doc, '1. 用户登录系统，进入"纹样基因库"页面', 12)
    add_paragraph(doc, '2. 使用搜索功能按主题、风格、地域检索纹样', 12)
    add_paragraph(doc, '3. 浏览纹样详情，查看版权信息', 12)
    add_paragraph(doc, '4. 免费纹样可直接下载使用', 12)
    add_paragraph(doc, '5. 版权纹样需购买授权后方可使用', 12)
    create_heading(doc, '4.4.3 运行截图', 3)
    add_paragraph(doc, '【此处插入基因库界面截图】', 12)
    
    create_heading(doc, '4.5 创意设计赋能业务', 2)
    create_heading(doc, '4.5.1 功能描述', 3)
    add_paragraph(doc, '提供在线设计白板与基因库素材，支持用户组合创新设计。涉及版权素材需购买授权后方可导出成果，支持发布设计"召集令"，面向平台用户招标征集作品。', 12)
    create_heading(doc, '4.5.2 操作步骤', 3)
    add_paragraph(doc, '1. 用户登录系统，进入"创意工坊"页面', 12)
    add_paragraph(doc, '2. 打开在线设计白板，选择基因库素材', 12)
    add_paragraph(doc, '3. 组合设计元素，创作原创作品', 12)
    add_paragraph(doc, '4. 使用版权素材时需购买授权', 12)
    add_paragraph(doc, '5. 完成设计后导出作品或发布"召集令"', 12)
    create_heading(doc, '4.5.3 运行截图', 3)
    add_paragraph(doc, '【此处插入创意工坊界面截图】', 12)
    
    create_heading(doc, '4.6 系统设置模块', 2)
    create_heading(doc, '4.6.1 功能描述', 3)
    add_paragraph(doc, '用户可修改个人信息（姓名、头像、联系方式）和登录密码，管理员可配置系统参数（如系统名称、版权信息等）。', 12)
    create_heading(doc, '4.6.2 操作步骤', 3)
    add_paragraph(doc, '1. 点击页面右上角的"个人中心"按钮', 12)
    add_paragraph(doc, '2. 进入"个人信息"页面，修改个人信息，点击"保存"按钮', 12)
    add_paragraph(doc, '3. 进入"修改密码"页面，输入原密码和新密码，点击"确认"按钮', 12)
    add_paragraph(doc, '4. 管理员进入"系统设置"页面，配置系统参数，点击"保存"按钮', 12)
    create_heading(doc, '4.6.3 运行截图', 3)
    add_paragraph(doc, '【此处插入系统设置界面截图】', 12)
    
    create_heading(doc, '第 5 章 异常处理说明', 1)
    add_paragraph(doc, '本软件对常见的异常情况进行了处理，确保系统稳定运行：', 12)
    add_paragraph(doc, '1. 网络异常：当网络连接中断时，系统会提示用户检查网络连接，并在网络恢复后自动重试操作', 12)
    add_paragraph(doc, '2. 认证失败：当用户登录凭证过期或无效时，系统自动跳转至登录页面，并提示用户重新登录', 12)
    add_paragraph(doc, '3. 权限不足：当用户尝试访问未授权的功能时，系统提示"没有权限执行此操作"', 12)
    add_paragraph(doc, '4. 参数错误：当用户提交的数据不符合要求时，系统提示具体的错误信息，如"用户名不能为空"', 12)
    add_paragraph(doc, '5. 资源不存在：当请求的资源不存在时，系统提示"请求的资源不存在"', 12)
    add_paragraph(doc, '6. 交易异常：当交易过程中出现余额不足、库存不足等情况时，系统提示相应的错误信息并取消交易', 12)
    add_paragraph(doc, '7. 文件上传失败：当文件上传失败时，系统提示"文件上传失败，请重试"', 12)
    add_paragraph(doc, '8. 服务器错误：当服务器内部出现错误时，系统提示"服务器内部错误，请稍后重试"', 12)
    
    create_heading(doc, '第 6 章 软件更新日志', 1)
    add_paragraph(doc, 'V1.0（2026年03月01日）：首次发布，实现用户登录、用户管理、数字版权交易业务、非遗基因库服务业务、创意设计赋能业务、系统设置六大功能模块。', 12)
    
    output_path = r'c:\Users\dell\Desktop\科创中心\大创\ZaRan-main\ZaRan-main\檐下风铃-扎染文化数字平台_V1.0_软件说明书.docx'
    doc.save(output_path)
    print(f"说明书已保存至: {output_path}")

if __name__ == "__main__":
    generate_manual()